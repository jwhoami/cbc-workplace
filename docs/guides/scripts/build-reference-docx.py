#!/usr/bin/env python3
"""
Genera docs/guides/templates/cbc-reference.docx — la plantilla Pandoc usada
por `pandoc --reference-doc=` para producir las tres guías de CBC Workplace.

La fuente de verdad de tokens (colores, tamaños, fuentes) vive en este
archivo. Para regenerar la plantilla:

    python3 docs/guides/scripts/build-reference-docx.py

Requiere: python-docx==1.1.2
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Cm, Twips
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx no instalado.\n"
        "Instalar con: pip3 install --user python-docx==1.1.2\n"
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Tokens de diseño — modificar aquí los valores cambia toda la plantilla
# ---------------------------------------------------------------------------

TOKENS = {
    # Colores (sin '#')
    "charcoal": "1F2937",
    "blue":     "2563EB",
    "bg":       "F9FAFB",
    "border":   "E5E7EB",
    "success":  "059669",
    "warn":     "D97706",
    "danger":   "DC2626",
    "muted":    "6B7280",
    "bg_note":    "EFF6FF",
    "bg_warn":    "FEF3C7",
    "bg_danger":  "FEE2E2",
    "bg_success": "D1FAE5",
    "bg_admin":   "F3F4F6",
    "white":    "FFFFFF",

    # Tipografías (deben estar instaladas en el sistema de quien genera/lee)
    "font_body":  "Inter",
    "font_mono":  "JetBrains Mono",

    # Tamaños en puntos (multiplicar por 2 para half-points en OOXML)
    "size_title":    36,
    "size_subtitle": 18,
    "size_h1":       24,
    "size_h2":       18,
    "size_h3":       14,
    "size_h4":       12,
    "size_body":     11,
    "size_caption":   9,
    "size_code":     10,
    "size_callout":  10.5,
}


# ---------------------------------------------------------------------------
# Helpers OOXML
# ---------------------------------------------------------------------------

def set_paragraph_shading(paragraph_format_or_style, hex_color):
    """Aplica un sombreado de fondo a un estilo de párrafo."""
    style_element = paragraph_format_or_style.element
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_paragraph_left_border(style, hex_color, size_pt=4):
    """Aplica un borde izquierdo grueso al estilo (para callouts)."""
    style_element = style.element
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt * 8))  # 1/8 of a point
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)


def set_paragraph_indent(style, left_cm=0.5, right_cm=0.3):
    """Indentación interna del párrafo (para callouts)."""
    style_element = style.element
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # 1 cm = 567 twips
    ind.set(qn("w:left"), str(int(left_cm * 567)))
    ind.set(qn("w:right"), str(int(right_cm * 567)))


# ---------------------------------------------------------------------------
# Constructor de la plantilla
# ---------------------------------------------------------------------------

def build_reference_doc(output_path: Path) -> None:
    doc = Document()

    # ---- Márgenes y orientación ------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    # ---- Estilo base "Normal" / Body Text -------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = TOKENS["font_body"]
    normal.font.size = Pt(TOKENS["size_body"])
    normal.font.color.rgb = RGBColor.from_string(TOKENS["charcoal"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.45

    # ---- Encabezados ------------------------------------------------------
    heading_specs = [
        ("Heading 1", TOKENS["size_h1"], True,  TOKENS["charcoal"], Pt(24), Pt(12)),
        ("Heading 2", TOKENS["size_h2"], True,  TOKENS["charcoal"], Pt(18), Pt(9)),
        ("Heading 3", TOKENS["size_h3"], True,  TOKENS["charcoal"], Pt(14), Pt(7)),
        ("Heading 4", TOKENS["size_h4"], True,  TOKENS["charcoal"], Pt(12), Pt(6)),
    ]
    for name, size_pt, bold, color_hex, before, after in heading_specs:
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = TOKENS["font_body"]
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color_hex)
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.keep_with_next = True

    # ---- Title / Subtitle ------------------------------------------------
    try:
        title_style = doc.styles["Title"]
    except KeyError:
        title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = TOKENS["font_body"]
    title_style.font.size = Pt(TOKENS["size_title"])
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(TOKENS["charcoal"])
    title_style.paragraph_format.space_after = Pt(12)

    try:
        subtitle_style = doc.styles["Subtitle"]
    except KeyError:
        subtitle_style = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = TOKENS["font_body"]
    subtitle_style.font.size = Pt(TOKENS["size_subtitle"])
    subtitle_style.font.color.rgb = RGBColor.from_string(TOKENS["muted"])
    subtitle_style.paragraph_format.space_after = Pt(24)

    # ---- Code (bloque) ---------------------------------------------------
    # Pandoc usa el estilo "Source Code" para code blocks
    try:
        code_block = doc.styles["Source Code"]
    except KeyError:
        code_block = doc.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    code_block.font.name = TOKENS["font_mono"]
    code_block.font.size = Pt(TOKENS["size_code"])
    code_block.font.color.rgb = RGBColor.from_string(TOKENS["charcoal"])
    code_block.paragraph_format.space_before = Pt(8)
    code_block.paragraph_format.space_after = Pt(8)
    code_block.paragraph_format.line_spacing = 1.3
    set_paragraph_shading(code_block, TOKENS["bg"])
    set_paragraph_indent(code_block, left_cm=0.3, right_cm=0.3)

    # ---- Code inline (carácter) -----------------------------------------
    try:
        code_char = doc.styles["Verbatim Char"]
    except KeyError:
        code_char = doc.styles.add_style("Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    code_char.font.name = TOKENS["font_mono"]
    code_char.font.size = Pt(TOKENS["size_body"] - 1)
    code_char.font.color.rgb = RGBColor.from_string(TOKENS["danger"])

    # ---- Caption ---------------------------------------------------------
    try:
        caption_style = doc.styles["Caption"]
    except KeyError:
        caption_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = TOKENS["font_body"]
    caption_style.font.size = Pt(TOKENS["size_caption"])
    caption_style.font.italic = True
    caption_style.font.color.rgb = RGBColor.from_string(TOKENS["muted"])
    caption_style.paragraph_format.space_after = Pt(12)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Quote -----------------------------------------------------------
    try:
        quote_style = doc.styles["Quote"]
    except KeyError:
        quote_style = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = TOKENS["font_body"]
    quote_style.font.size = Pt(TOKENS["size_body"])
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor.from_string(TOKENS["muted"])
    set_paragraph_indent(quote_style, left_cm=0.8, right_cm=0.8)
    set_paragraph_left_border(quote_style, TOKENS["border"], size_pt=2)

    # ---- Callouts --------------------------------------------------------
    callouts = [
        ("Callout-Note",      TOKENS["blue"],    TOKENS["bg_note"]),
        ("Callout-Warn",      TOKENS["warn"],    TOKENS["bg_warn"]),
        ("Callout-Danger",    TOKENS["danger"],  TOKENS["bg_danger"]),
        ("Callout-Success",   TOKENS["success"], TOKENS["bg_success"]),
        ("Callout-AdminOnly", TOKENS["charcoal"],TOKENS["bg_admin"]),
    ]
    for name, border_color, bg_color in callouts:
        try:
            cstyle = doc.styles[name]
        except KeyError:
            cstyle = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        cstyle.font.name = TOKENS["font_body"]
        cstyle.font.size = Pt(TOKENS["size_callout"])
        cstyle.font.color.rgb = RGBColor.from_string(TOKENS["charcoal"])
        cstyle.paragraph_format.space_before = Pt(8)
        cstyle.paragraph_format.space_after = Pt(8)
        cstyle.paragraph_format.line_spacing = 1.35
        set_paragraph_shading(cstyle, bg_color)
        set_paragraph_left_border(cstyle, border_color, size_pt=3)
        set_paragraph_indent(cstyle, left_cm=0.4, right_cm=0.4)

    # ---- Estilos de tabla (Pandoc usa "Table" por defecto) ---------------
    # Pandoc consume el estilo "Table Caption" + el estilo de tabla nombrado
    # en el reference doc; mantener el default es suficiente.

    # ---- Contenido de ejemplo (visible solo si Pandoc no lo sobreescribe)
    # Pandoc reemplaza el contenido — el body del template sólo importa
    # para los estilos. Igual añadimos un párrafo de marca para humanos que
    # abran este .docx directo.
    doc.add_paragraph(
        "Esta es la plantilla de referencia de CBC Workplace. "
        "No edite directamente: regenere con "
        "python3 docs/guides/scripts/build-reference-docx.py.",
        style="Subtitle",
    )

    # ---- Footer con versión ----------------------------------------------
    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.text = "CBC Workplace · Guías oficiales · v1.0 — Mayo 2026"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(TOKENS["muted"])
            run.font.name = TOKENS["font_body"]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[ok] Plantilla generada en: {output_path}")
    print(f"[ok] Tamaño: {output_path.stat().st_size} bytes")


def main() -> int:
    repo_root = Path(__file__).resolve().parents[3]
    output_path = repo_root / "docs" / "guides" / "templates" / "cbc-reference.docx"
    build_reference_doc(output_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
